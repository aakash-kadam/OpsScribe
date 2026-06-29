"""DOCX writer for generated OpsScribe reports."""

from __future__ import annotations

import re
from pathlib import Path
from collections.abc import Sequence

from docx import Document
from docx.shared import Inches
from docx.text.paragraph import Paragraph
from docx.shared import Pt

from report_models import ReportResult


def is_markdown_table_row(line: str) -> bool:
    """Return whether a line looks like a markdown table row."""
    stripped = line.strip()
    return stripped.startswith("|") and stripped.endswith("|") and "|" in stripped[1:-1]


def is_markdown_table_separator(line: str) -> bool:
    """Return whether a markdown table row is the alignment separator."""
    cells = parse_table_row(line)
    return bool(cells) and all(re.fullmatch(r":?-{3,}:?", cell.strip()) for cell in cells)


def parse_table_row(line: str) -> list[str]:
    """Parse a markdown table row into cell values."""
    return [cell.strip() for cell in line.strip().strip("|").split("|")]


def add_bold_runs(paragraph: Paragraph, text: str) -> None:
    """Add text to a paragraph, treating **text** spans as bold."""
    parts = re.split(r"(\*\*.*?\*\*)", text)
    for part in parts:
        if not part:
            continue

        if part.startswith("**") and part.endswith("**") and len(part) >= 4:
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        else:
            paragraph.add_run(part)


def add_paragraph_with_bold(document: Document, text: str, style: str | None = None) -> None:
    """Add a paragraph while preserving markdown-style bold spans."""
    paragraph = document.add_paragraph(style=style)
    add_bold_runs(paragraph, text)


def add_markdown_table(document: Document, table_lines: list[str]) -> None:
    """Add a markdown table as a real DOCX table."""
    rows = [parse_table_row(line) for line in table_lines if not is_markdown_table_separator(line)]
    if not rows:
        return

    column_count = max(len(row) for row in rows)
    table = document.add_table(rows=len(rows), cols=column_count)
    table.style = "Table Grid"

    for row_index, row in enumerate(rows):
        for column_index in range(column_count):
            cell_text = row[column_index] if column_index < len(row) else ""
            paragraph = table.cell(row_index, column_index).paragraphs[0]
            add_bold_runs(paragraph, cell_text)

    document.add_paragraph()


def add_markdownish_content(document: Document, content: str) -> None:
    """Add basic markdown-like text to a Word document."""
    in_code_block = False
    lines = content.splitlines()
    index = 0

    while index < len(lines):
        raw_line = lines[index]
        line = raw_line.strip()

        if not line:
            index += 1
            continue
        if line.startswith("```"):
            in_code_block = not in_code_block
            index += 1
            continue
        if in_code_block:
            paragraph = document.add_paragraph(line)
            for run in paragraph.runs:
                run.font.name = "Courier New"
                run.font.size = Pt(9)
            index += 1
            continue

        if is_markdown_table_row(line):
            table_lines = []
            while index < len(lines) and is_markdown_table_row(lines[index].strip()):
                table_lines.append(lines[index].strip())
                index += 1
            add_markdown_table(document, table_lines)
            continue

        if line.startswith("### "):
            add_bold_runs(document.add_heading(level=3), line.removeprefix("### "))
        elif line.startswith("## "):
            add_bold_runs(document.add_heading(level=2), line.removeprefix("## "))
        elif line.startswith("# "):
            add_bold_runs(document.add_heading(level=1), line.removeprefix("# "))
        elif line.startswith("- "):
            add_paragraph_with_bold(document, line.removeprefix("- "), style="List Bullet")
        elif re.match(r"^\d+\.\s+", line):
            _, item = line.split(". ", 1)
            add_paragraph_with_bold(document, item, style="List Number")
        else:
            add_paragraph_with_bold(document, line)

        index += 1


def add_graphs(document: Document, graph_paths: Sequence[Path]) -> None:
    """Add graph images to the document under a visual summary section."""
    if not graph_paths:
        return

    document.add_heading("Visual Case summarization", level=1)
    for graph_path in graph_paths:
        title = graph_path.stem.replace("_", " ").title()
        document.add_paragraph(title)
        document.add_picture(str(graph_path), width=Inches(6.5))


def write_report_docx(report: ReportResult, output_path: Path, graph_paths: Sequence[Path] | None = None) -> None:
    """Write a generated report to a DOCX file."""
    document = Document()
    document.add_heading(report.title, level=0)
    add_graphs(document, graph_paths or [])

    for section in report.sections:
        document.add_heading(section.title, level=1)
        add_markdownish_content(document, section.content)

    usage = report.token_usage
    if usage.available:
        document.add_heading("Token Usage", level=1)
        document.add_paragraph(f"Input tokens: {usage.input_tokens}")
        document.add_paragraph(f"Output tokens: {usage.output_tokens}")
        document.add_paragraph(f"Total tokens: {usage.total_tokens}")

    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(output_path)
