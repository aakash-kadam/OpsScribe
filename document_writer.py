"""Document writers for generated OpsScribe reports."""

from __future__ import annotations

from pathlib import Path

from docx import Document
from reportlab.lib.pagesizes import LETTER
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer


def write_markdown_report(markdown: str, output_path: Path) -> None:
    """Write markdown text to disk."""
    output_path.write_text(markdown, encoding="utf-8")


def write_docx_report(markdown: str, output_path: Path) -> None:
    """Write a basic Word document from markdown-like report text."""
    document = Document()

    for line in markdown.splitlines():
        line = line.strip()

        if not line:
            continue
        if line.startswith("# "):
            document.add_heading(line.removeprefix("# "), level=0)
        elif line.startswith("## "):
            document.add_heading(line.removeprefix("## "), level=1)
        elif line.startswith("### "):
            document.add_heading(line.removeprefix("### "), level=2)
        elif line.startswith("- "):
            document.add_paragraph(line.removeprefix("- "), style="List Bullet")
        elif line.startswith("```"):
            continue
        else:
            document.add_paragraph(line)

    document.save(output_path)


def write_pdf_report(markdown: str, output_path: Path) -> None:
    """Write a basic PDF document from markdown-like report text."""
    styles = getSampleStyleSheet()
    document = SimpleDocTemplate(str(output_path), pagesize=LETTER)
    elements = []

    for line in markdown.splitlines():
        line = line.strip()

        if not line:
            elements.append(Spacer(1, 8))
        elif line.startswith("# "):
            elements.append(Paragraph(line.removeprefix("# "), styles["Title"]))
        elif line.startswith("## "):
            elements.append(Paragraph(line.removeprefix("## "), styles["Heading1"]))
        elif line.startswith("### "):
            elements.append(Paragraph(line.removeprefix("### "), styles["Heading2"]))
        elif line.startswith("- "):
            elements.append(Paragraph(f"- {line.removeprefix('- ')}", styles["BodyText"]))
        elif line.startswith("```"):
            continue
        else:
            elements.append(Paragraph(line, styles["BodyText"]))

    document.build(elements)
